#!/usr/bin/env python3
"""
Vayva Marketing Plan Document Generator v4
Corrected pricing - NO Free forever plan (Starter and Pro only, with 7-day trial)
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import cairosvg
import os
from PIL import Image
import io

# Create document
doc = Document()

# Set up styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Page setup
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Convert SVG logo to PNG for header
logo_svg_path = '/Users/fredrick/Documents/Vayva-Tech/vayva/Frontend/marketing/public/vayva-logo-official.svg'
logo_png_path = '/Users/fredrick/Documents/Vayva-Tech/vayva/vayva_logo_header.png'

# Convert SVG to PNG with proper sizing
if os.path.exists(logo_svg_path):
    png_data = cairosvg.svg2png(url=logo_svg_path, output_width=120, output_height=40)
    with open(logo_png_path, 'wb') as f:
        f.write(png_data)

# Add header with logo on LEFT and Vayva text on RIGHT
def add_header(section):
    header = section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.clear()

    # Create table for header layout (logo left, text right)
    table = header.add_table(1, 2, Inches(6.5))
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(1.2)
    table.columns[1].width = Inches(5.3)

    # Remove table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    # Cell 1: Logo
    cell_logo = table.cell(0, 0)
    cell_logo_para = cell_logo.paragraphs[0]
    cell_logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if os.path.exists(logo_png_path):
        run = cell_logo_para.add_run()
        run.add_picture(logo_png_path, width=Inches(1.0))

    # Cell 2: Vayva text
    cell_text = table.cell(0, 1)
    cell_text_para = cell_text.paragraphs[0]
    cell_text_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = cell_text_para.add_run("Vayva")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Add separator line
    header.add_paragraph("_" * 80)

# Apply header to first section
add_header(sections[0])

# Helper functions
def add_heading1(text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_heading2(text):
    p = doc.add_heading(text, level=2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_heading3(text):
    p = doc.add_heading(text, level=3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_paragraph(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet' if level == 0 else 'List Bullet 2')
    return p

def add_numbered(text):
    p = doc.add_paragraph(text, style='List Number')
    return p

# ==================== DOCUMENT CONTENT ====================

# Title
add_heading1("VAYVA MARKETING STRATEGY & GROWTH PLAYBOOK")
add_paragraph("The Complete Guide to Dominating African E-Commerce Through Conversational Commerce", italic=True)
add_paragraph("Version 4.0 - March 2026")
doc.add_paragraph()

# Executive Summary
add_heading1("EXECUTIVE SUMMARY")
add_paragraph("Vayva is the commerce operating system built specifically for African businesses. Unlike generic e-commerce platforms, Vayva transforms WhatsApp and Instagram conversations into structured business operations—automating order capture, inventory management, payments, and delivery tracking.")
doc.add_paragraph()

add_heading2("Core Value Proposition")
add_bullet("AI-powered order capture from WhatsApp and Instagram conversations")
add_bullet("Paystack-integrated payments (cards, transfers, USSD, mobile money)")
add_bullet("Kwik Delivery integration for logistics management")
add_bullet("Voice commerce in Nigerian languages (Yoruba, Igbo, Hausa, Pidgin)")
add_bullet("Buy Now Pay Later (BNPL) functionality for customer flexibility")
add_bullet("A/B testing framework for conversion optimization")
add_bullet("Public checkout API for embedding commerce anywhere")
doc.add_paragraph()

add_heading2("Corrected Pricing Structure")
add_paragraph("Vayva offers two paid plans with a 7-day free trial:", bold=True)
doc.add_paragraph()

# Pricing table
pricing_table = doc.add_table(rows=3, cols=3)
pricing_table.style = 'Light Grid Accent 1'
pricing_table.cell(0, 0).text = "Plan"
pricing_table.cell(0, 1).text = "Monthly Price"
pricing_table.cell(0, 2).text = "Key Features"

pricing_table.cell(1, 0).text = "Starter"
pricing_table.cell(1, 1).text = "₦25,000/month"
pricing_table.cell(1, 2).text = "Up to 500 products, WhatsApp & Instagram automation, Storefront setup, Remove Vayva branding, 7-day free trial"

pricing_table.cell(2, 0).text = "Pro"
pricing_table.cell(2, 1).text = "₦40,000/month"
pricing_table.cell(2, 2).text = "Everything in Starter + Multi-location support, API access, Dedicated account manager, Custom integrations"

doc.add_paragraph()
add_paragraph("Note: 7-day free trial available for Starter plan. No credit card required. After trial, account is paused if not subscribed.", italic=True)
doc.add_paragraph()

# Target Market
add_heading1("TARGET MARKET ANALYSIS")
add_heading2("Primary Segments")
add_bullet("Instagram vendors and WhatsApp sellers (side hustles to full-time)")
add_bullet("Small retail businesses with physical inventory")
add_bullet("Growing brands needing automation and team collaboration")
add_bullet("Fashion, beauty, electronics, grocery, and service-based businesses")
doc.add_paragraph()

add_heading2("Market Size (Nigeria)")
add_bullet("40+ million SMEs in Nigeria")
add_bullet("90% of online commerce happens via WhatsApp and Instagram")
add_bullet("$8+ billion informal e-commerce market")
add_bullet("Average merchant processes 50-200 orders/month manually")
doc.add_paragraph()

add_heading2("Pain Points We Solve")
add_numbered("Manual order tracking via screenshots and spreadsheets")
add_numbered("Lost sales from missed WhatsApp messages")
add_numbered("Payment confirmation delays and fraud")
add_numbered("Inventory stockouts and overselling")
add_numbered("Delivery coordination chaos")
add_numbered("No analytics on what's selling")
doc.add_paragraph()

# Platform-by-Platform Playbook
add_heading1("PLATFORM-BY-PLATFORM MARKETING PLAYBOOK")
add_paragraph("Exact step-by-step instructions for dominating each platform at low cost.")
doc.add_paragraph()

# Twitter/X Strategy
add_heading2("1. TWITTER/X STRATEGY")
add_heading3("The Nigerian Business Twitter Formula")
doc.add_paragraph()

add_paragraph("Daily Content Framework:", bold=True)
add_numbered("Post at 8-9 AM (commute time) - Business motivation/insights")
add_numbered("Post at 12-1 PM (lunch break) - Educational thread")
add_numbered("Post at 6-8 PM (evening scroll) - Customer success story")
doc.add_paragraph()

add_paragraph("Viral Tweet Templates:", bold=True)
add_bullet("'I helped [X] vendors automate [Y] orders last month. Here's what I learned...'")
add_bullet("'Stop doing [manual task]. Here's the automation hack...'")
add_bullet("'Before Vayva vs After Vayva' screenshot threads")
add_bullet("'The real reason your WhatsApp business isn't scaling...'")
doc.add_paragraph()

add_paragraph("Growth Hacking Tactics:", bold=True)
add_bullet("Reply to every business-related tweet with value-first comments")
add_bullet("Quote tweet industry news with Vayva angle")
add_bullet("Create 'Vendors of Lagos/Abuja/PH' appreciation threads")
add_bullet("Run polls: 'What's your biggest order management struggle?'")
add_bullet("Use Spaces to host 'Vendor Office Hours' weekly")
doc.add_paragraph()

add_paragraph("Hashtag Strategy:", bold=True)
add_bullet("Primary: #Vayva #NigerianBusiness #WhatsAppCommerce")
add_bullet("Secondary: #VendorLife #SmallBusinessNigeria #EcommerceTips")
add_bullet("Trending: Jump on relevant Nigerian business hashtags")
doc.add_paragraph()

# LinkedIn Strategy
add_heading2("2. LINKEDIN STRATEGY")
add_heading3("B2B and Professional Positioning")
doc.add_paragraph()

add_paragraph("Content Pillars:", bold=True)
add_numbered("African commerce insights and market trends")
add_numbered("Founder journey and building in public")
add_numbered("Customer case studies and success stories")
add_numbered("E-commerce education and best practices")
doc.add_paragraph()

add_paragraph("Posting Schedule:", bold=True)
add_bullet("Tuesday 8 AM: Long-form article (1500+ words)")
add_bullet("Thursday 12 PM: Carousel post (5-7 slides)")
add_bullet("Friday 4 PM: Personal founder story/reflection")
doc.add_paragraph()

add_paragraph("LinkedIn-Specific Tactics:", bold=True)
add_bullet("Tag relevant connections in thoughtful comments")
add_bullet("Engage with every comment within 1 hour")
add_bullet("Use LinkedIn newsletters for weekly 'Commerce Roundup'")
add_bullet("Connect with 20 new business owners daily")
add_bullet("Join and contribute to Nigerian business groups")
doc.add_paragraph()

# Instagram Strategy
add_heading2("3. INSTAGRAM STRATEGY")
add_heading3("Visual Commerce and Vendor Community")
doc.add_paragraph()

add_paragraph("Content Mix (Weekly):", bold=True)
add_bullet("4 Reels: Quick tips, customer transformations, behind-the-scenes")
add_bullet("3 Carousels: Educational content, step-by-step guides")
add_bullet("5 Stories: Daily engagement, polls, Q&A, tips")
add_bullet("1 Live: Weekly 'Vendor Office Hours'")
doc.add_paragraph()

add_paragraph("Reel Ideas That Convert:", bold=True)
add_numbered("'POV: You discovered Vayva after 6 months of spreadsheet hell'")
add_numbered("'How [specific vendor] processes 100 orders in 10 minutes'")
add_numbered("'The automation your competitors don't want you to know'")
add_numbered("'Before vs After: Order management edition'")
add_numbered("'Day in the life of a Vayva-powered vendor'")
doc.add_paragraph()

add_paragraph("Instagram Growth Hacks:", bold=True)
add_bullet("Comment on 50 vendor posts daily with genuine value")
add_bullet("Use location tags: Lagos, Abuja, Port Harcourt, Ibadan")
add_bullet("Collaborate with vendor influencers for takeovers")
add_bullet("Create 'Vendor Spotlight' series")
add_bullet("Run 'Comment VAYVA for a DM' engagement posts")
doc.add_paragraph()

# TikTok Strategy
add_heading2("4. TIKTOK STRATEGY")
add_heading3("Viral Commerce Education")
doc.add_paragraph()

add_paragraph("Video Formats That Work:", bold=True)
add_bullet("'Things I wish I knew before starting my online business'")
add_bullet("'How I automated my WhatsApp orders (tutorial)'")
add_bullet("'Small business glow up - manual to automated'")
add_bullet("'Nigerian vendor problems that need solutions'")
add_bullet("'Testing if Vayva actually works - Day 1 to Day 30'")
doc.add_paragraph()

add_paragraph("TikTok Growth Formula:", bold=True)
add_numbered("Post 2-3x daily (algorithm favors consistency)")
add_numbered("Use trending sounds with business twist")
add_numbered("Hook in first 3 seconds: 'Stop scrolling if you sell on WhatsApp...'")
add_numbered("Include text overlays for sound-off viewing")
add_numbered("End with clear CTA: 'Link in bio to automate your orders'")
doc.add_paragraph()

# WhatsApp Strategy
add_heading2("5. WHATSAPP STRATEGY")
add_heading3("The Primary Channel (This is Vayva's Home)")
doc.add_paragraph()

add_paragraph("WhatsApp Business Profile Optimization:", bold=True)
add_bullet("Profile: 'Vayva - Turn WhatsApp chats into organized sales'")
add_bullet("Catalog: Feature 4 key use cases as 'products'")
add_bullet("Status: Daily tips, customer wins, feature highlights")
add_bullet("About: Link to signup + '7-day free trial available'")
doc.add_paragraph()

add_paragraph("WhatsApp Status Content Calendar:", bold=True)
add_bullet("Monday: Motivation + feature highlight")
add_bullet("Tuesday: Customer testimonial")
add_bullet("Wednesday: Tutorial/How-to")
add_bullet("Thursday: Before/After transformation")
add_bullet("Friday: Weekend business tip")
add_bullet("Saturday: Community spotlight")
add_bullet("Sunday: Week ahead preparation")
doc.add_paragraph()

add_paragraph("WhatsApp Broadcast Strategy:", bold=True)
add_numbered("Segment leads by industry (fashion, beauty, electronics)")
add_numbered("Send weekly value-packed broadcast (not just sales)")
add_numbered("Personalize with first names")
add_numbered("Include 'Reply STOP to unsubscribe'")
add_numbered("Track opens via link clicks")
doc.add_paragraph()

# Nairaland Strategy
add_heading2("6. NAIRALAND STRATEGY")
add_heading3("Nigeria's Largest Forum (Undervalued)")
doc.add_paragraph()

add_paragraph("Target Boards:", bold=True)
add_bullet("Business")
add_bullet("E-commerce")
add_bullet("Technology")
add_bullet("Career")
add_bullet("Investment")
doc.add_paragraph()

add_paragraph("Nairaland Tactics:", bold=True)
add_numbered("Create valuable threads: 'How to scale your WhatsApp business'")
add_numbered("Answer questions genuinely - signature links to Vayva")
add_numbered("Share case studies with real numbers")
add_numbered("Position as helpful expert, not salesperson")
add_numbered("Post during high-traffic hours: 8-10 AM, 6-9 PM")
doc.add_paragraph()

add_paragraph("Thread Templates:", bold=True)
add_bullet("'I automated my wife's hair business - here's what happened'")
add_bullet("'The real cost of manual order management (calculations inside)'")
add_bullet("'Tools every Nigerian online vendor should know about'")
doc.add_paragraph()

# Reddit Strategy
add_heading2("7. REDDIT STRATEGY")
add_heading3("Authentic Community Engagement")
doc.add_paragraph()

add_paragraph("Target Subreddits:", bold=True)
add_bullet("r/Nigeria")
add_bullet("r/AfricanBusiness")
add_bullet("r/ecommerce")
add_bullet("r/smallbusiness")
add_bullet("r/entrepreneur")
doc.add_paragraph()

add_paragraph("Reddit Rules:", bold=True)
add_numbered("Provide 10x value before any mention of Vayva")
add_numbered("Share real data and experiences")
add_numbered("Be transparent about affiliation")
add_numbered("Answer DMs promptly")
add_numbered("Build karma through genuine contributions")
doc.add_paragraph()

# Product Hunt Strategy
add_heading2("8. PRODUCT HUNT STRATEGY")
add_heading3("Tech Community Launch")
doc.add_paragraph()

add_paragraph("Pre-Launch (2 weeks before):", bold=True)
add_bullet("Build email list of supporters")
add_bullet("Create teaser content")
add_bullet("Reach out to Nigerian tech community")
add_bullet("Prepare all assets (screenshots, video, description)")
doc.add_paragraph()

add_paragraph("Launch Day Execution:", bold=True)
add_numbered("Post at 12:01 AM PST (optimal for global reach)")
add_numbered("Share across all channels simultaneously")
add_numbered("Respond to every comment within minutes")
add_numbered("Reach out to personal network for upvotes")
add_numbered("Live-tweet the launch journey")
doc.add_paragraph()

add_paragraph("Post-Launch:", bold=True)
add_bullet("Thank every supporter personally")
add_bullet("Convert hunters to users with special offer")
add_bullet("Document lessons learned")
add_bullet("Build relationships with commenters")
doc.add_paragraph()

# Growth Hacking Tactics
add_heading1("GROWTH HACKING TACTICS")
add_paragraph("Secret methods used by successful companies to acquire users at low cost.")
doc.add_paragraph()

add_heading2("1. The 'Powered by Vayva' Viral Loop")
add_paragraph("Every store created on Vayva includes a small 'Powered by Vayva' badge in the footer. Clicking it leads to a landing page. Offer merchants 10% off their next month for keeping the badge visible.")
doc.add_paragraph()

add_heading2("2. Comment Hijacking")
add_paragraph("Find viral posts about business struggles on Twitter, LinkedIn, Instagram. Comment with genuine value + soft mention of Vayva. Focus on posts with high engagement but low-quality comments.")
doc.add_paragraph()

add_heading2("3. The 'Vendor Community' Strategy")
add_paragraph("Create a private WhatsApp or Telegram group for Vayva users. Share exclusive tips, early features, and networking opportunities. Members become evangelists.")
doc.add_paragraph()

add_heading2("4. Cross-Pollination Content")
add_paragraph("Turn one piece of content into 10:")
add_bullet("Twitter thread becomes LinkedIn article")
add_bullet("LinkedIn article becomes Instagram carousel")
add_bullet("Instagram carousel becomes TikTok slideshow")
add_bullet("TikTok becomes YouTube Short")
add_bullet("All become email newsletter content")
doc.add_paragraph()

add_heading2("5. The 'Free Tool' Lead Magnet")
add_paragraph("Create free utilities that attract target users:")
add_bullet("WhatsApp Order Template (Google Sheets)")
add_bullet("E-commerce Profit Calculator")
add_bullet("Inventory Management Checklist")
add_bullet("'50 Business Ideas for 2026' guide")
doc.add_paragraph()

add_heading2("6. Referral Program Mechanics")
add_paragraph("Structure: Referrer gets ₦5,000 credit, Referee gets 20% off first month. Make it stupidly easy to share with one-click WhatsApp sharing.")
doc.add_paragraph()

add_heading2("7. Strategic Partnerships")
add_paragraph("Partner with complementary services:")
add_bullet("Paystack (payment processor)")
add_bullet("Kwik Delivery (logistics)")
add_bullet("Instagram growth agencies")
add_bullet("Business coaches and consultants")
add_bullet("Packaging suppliers")
doc.add_paragraph()

# Paid Advertising
add_heading1("PAID ADVERTISING STRATEGY")
add_heading2("Meta Ads (Facebook & Instagram)")
doc.add_paragraph()

add_paragraph("Campaign Structure:", bold=True)
add_bullet("Awareness: Video views of customer success stories")
add_bullet("Consideration: Traffic to industry-specific landing pages")
add_bullet("Conversion: Free trial signups with retargeting")
doc.add_paragraph()

add_paragraph("Ad Creative Best Practices:", bold=True)
add_numbered("Use UGC (user-generated content) style videos")
add_numbered("Show real Nigerian vendors using the platform")
add_numbered("Lead with pain point, end with transformation")
add_numbered("Include '7-day free trial' in first 3 seconds")
add_numbered("Use Nigerian Pidgin for broader appeal")
doc.add_paragraph()

add_paragraph("Targeting:", bold=True)
add_bullet("Interests: E-commerce, Small business, Entrepreneurship")
add_bullet("Behaviors: Online shoppers, Engaged shoppers")
add_bullet("Custom: Website visitors, Email list uploads")
add_bullet("Lookalike: Based on existing customers")
doc.add_paragraph()

add_heading2("Google Ads")
doc.add_paragraph()

add_paragraph("Search Campaign Keywords:", bold=True)
add_bullet("High intent: 'WhatsApp business automation', 'order management system Nigeria'")
add_bullet("Problem-aware: 'how to track WhatsApp orders', 'inventory management for vendors'")
add_bullet("Competitor: Alternative searches")
doc.add_paragraph()

add_heading2("TikTok Ads")
doc.add_paragraph()
add_paragraph("Best for awareness and trial signups. Use Spark Ads to boost organic content. Target 'Small Business Owners' interest category.")
doc.add_paragraph()

# Email Marketing
add_heading1("EMAIL MARKETING AUTOMATION")
add_heading2("Welcome Sequence (7 emails over 14 days)")
doc.add_paragraph()

add_numbered("Day 0: Welcome + quick win setup guide")
add_numbered("Day 2: Success story from similar business")
add_numbered("Day 4: Feature deep-dive (WhatsApp automation)")
add_numbered("Day 6: Social proof (testimonials + case studies)")
add_numbered("Day 8: Address objections (pricing, complexity)")
add_numbered("Day 10: Urgency (trial ending soon)")
add_numbered("Day 14: Final chance + alternative options")
doc.add_paragraph()

add_heading2("Monthly Newsletter")
add_paragraph("Content mix: 40% educational, 30% product updates, 20% customer stories, 10% promotional")
doc.add_paragraph()

# SEO Strategy
add_heading1("SEO & CONTENT STRATEGY")
add_heading2("Target Keywords (Priority Order)")
doc.add_paragraph()

add_bullet("Primary: 'WhatsApp business Nigeria', 'order management system', 'e-commerce automation'")
add_bullet("Secondary: 'Instagram vendor tools', 'inventory management Nigeria', 'online business automation'")
add_bullet("Long-tail: 'how to manage WhatsApp orders', 'best platform for Instagram vendors'")
doc.add_paragraph()

add_heading2("Content Clusters")
add_paragraph("Create pillar pages with supporting content:")
add_bullet("Pillar: 'The Complete Guide to WhatsApp Commerce'")
add_bullet("Cluster: 'How to set up WhatsApp Business API', 'WhatsApp automation tools', 'WhatsApp marketing strategies'")
doc.add_paragraph()

add_heading2("Technical SEO")
add_bullet("Optimize for Core Web Vitals")
add_bullet("Implement schema markup for rich snippets")
add_bullet("Create XML sitemap and submit to Google")
add_bullet("Build internal linking structure")
doc.add_paragraph()

# PR & Media
add_heading1("PR & MEDIA STRATEGY")
add_heading2("Nigerian Tech Media Targets")
add_bullet("TechCabal - Pitch founder story or funding news")
add_bullet("Techpoint Africa - Product launch or feature")
add_bullet("Disrupt Africa - Pan-African angle")
add_bullet("BusinessDay - SME-focused story")
add_bullet("Nairametrics - Business model deep-dive")
doc.add_paragraph()

add_heading2("PR Angles That Work")
add_numbered("Founder story: 'Why I quit my banking job to build for vendors'")
add_numbered("Market insight: 'The $8B WhatsApp economy nobody talks about'")
add_numbered("Customer milestone: 'How Vayva helped 1,000 vendors automate'")
add_numbered("Innovation: 'First voice commerce in Nigerian languages'")
doc.add_paragraph()

# Metrics & KPIs
add_heading1("METRICS & KPIs")
add_heading2("North Star Metric")
add_paragraph("Monthly Active Merchants (MAM) - merchants processing >10 orders/month")
doc.add_paragraph()

add_heading2("Key Metrics to Track")
table = doc.add_table(rows=9, cols=3)
table.style = 'Light Grid Accent 1'
table.cell(0, 0).text = "Metric"
table.cell(0, 1).text = "Target"
table.cell(0, 2).text = "Measurement"

table.cell(1, 0).text = "Trial Signups"
table.cell(1, 1).text = "500/month"
table.cell(1, 2).text = "Website conversions"

table.cell(2, 0).text = "Trial-to-Paid Conversion"
table.cell(2, 1).text = "15-20%"
table.cell(2, 2).text = "Revenue tracking"

table.cell(3, 0).text = "CAC (Customer Acquisition Cost)"
table.cell(3, 1).text = "<₦5,000"
table.cell(3, 2).text = "Marketing spend / customers"

table.cell(4, 0).text = "LTV (Lifetime Value)"
table.cell(4, 1).text = ">₦150,000"
table.cell(4, 2).text = "Avg months * monthly price"

table.cell(5, 0).text = "Churn Rate"
table.cell(5, 1).text = "<5%/month"
table.cell(5, 2).text = "Canceled / total customers"

table.cell(6, 0).text = "NRR (Net Revenue Retention)"
table.cell(6, 1).text = ">100%"
table.cell(6, 2).text = "Expansion - contraction"

table.cell(7, 0).text = "Social Followers"
table.cell(7, 1).text = "10,000/quarter"
table.cell(7, 2).text = "Platform analytics"

table.cell(8, 0).text = "Organic Traffic"
table.cell(8, 1).text = "50% MoM growth"
table.cell(8, 2).text = "Google Analytics"

doc.add_paragraph()

# 90-Day Action Plan
add_heading1("90-DAY ACTION PLAN")
add_heading2("Month 1: Foundation")
add_bullet("Set up all social media profiles with consistent branding")
add_bullet("Create content calendar and batch-produce 30 days of content")
add_bullet("Launch referral program")
add_bullet("Begin daily engagement on Twitter, Instagram, LinkedIn")
add_bullet("Set up email automation sequences")
add_bullet("Implement analytics tracking")
doc.add_paragraph()

add_heading2("Month 2: Acceleration")
add_bullet("Launch paid ads on Meta (₦200,000 budget)")
add_bullet("Publish 4 SEO-optimized blog posts")
add_bullet("Host first 'Vendor Office Hours' webinar")
add_bullet("Reach out to 10 potential strategic partners")
add_bullet("Create and distribute first lead magnet")
add_bullet("Begin Nairaland engagement strategy")
doc.add_paragraph()

add_heading2("Month 3: Scale")
add_bullet("Launch Product Hunt campaign")
add_bullet("Pitch to 5 Nigerian tech journalists")
add_bullet("Implement 'Powered by Vayva' viral loop")
add_bullet("Launch vendor community group")
add_bullet("Analyze metrics and double down on top 3 channels")
add_bullet("Plan next quarter based on learnings")
doc.add_paragraph()

# Budget Allocation
add_heading1("BUDGET ALLOCATION (MONTHLY)")
table = doc.add_table(rows=9, cols=3)
table.style = 'Light Grid Accent 1'
table.cell(0, 0).text = "Channel"
table.cell(0, 1).text = "Budget (₦)"
table.cell(0, 2).text = "Expected Outcome"

table.cell(1, 0).text = "Meta Ads (FB/IG)"
table.cell(1, 1).text = "150,000"
table.cell(1, 2).text = "100 trial signups"

table.cell(2, 0).text = "Google Ads"
table.cell(2, 1).text = "50,000"
table.cell(2, 2).text = "30 high-intent signups"

table.cell(3, 0).text = "TikTok Ads"
table.cell(3, 1).text = "50,000"
table.cell(3, 2).text = "50 signups, brand awareness"

table.cell(4, 0).text = "Content Creation"
table.cell(4, 1).text = "75,000"
table.cell(4, 2).text = "30 posts, 10 videos"

table.cell(5, 0).text = "Influencer Partnerships"
table.cell(5, 1).text = "50,000"
table.cell(5, 2).text = "5 micro-influencer posts"

table.cell(6, 0).text = "Tools & Software"
table.cell(6, 1).text = "25,000"
table.cell(6, 2).text = "Analytics, design, automation"

table.cell(7, 0).text = "Events/Community"
table.cell(7, 1).text = "50,000"
table.cell(7, 2).text = "1 event, community building"

table.cell(8, 0).text = "TOTAL"
table.cell(8, 1).text = "450,000"
table.cell(8, 2).text = "180+ trial signups"

doc.add_paragraph()

# Competitive Positioning
add_heading1("COMPETITIVE POSITIONING")
add_heading2("Vayva vs. Competitors")
doc.add_paragraph()

add_paragraph("vs. Shopify:", bold=True)
add_bullet("Vayva: Built for WhatsApp-first African commerce")
add_bullet("Shopify: Generic e-commerce, expensive for African market")
doc.add_paragraph()

add_paragraph("vs. Selar:", bold=True)
add_bullet("Vayva: AI automation + inventory + delivery integration")
add_bullet("Selar: Digital products focus, limited automation")
doc.add_paragraph()

add_paragraph("vs. Manual WhatsApp:", bold=True)
add_bullet("Vayva: Organized, automated, scalable")
add_bullet("Manual: Chaotic, time-consuming, error-prone")
doc.add_paragraph()

add_paragraph("vs. Flutterwave Commerce:", bold=True)
add_bullet("Vayva: Conversation-first, AI-powered")
add_bullet("Flutterwave: Payment-first, less automation")
doc.add_paragraph()

# Conclusion
add_heading1("CONCLUSION")
add_paragraph("Vayva is positioned to capture the massive informal e-commerce market in Nigeria and across Africa. By focusing on WhatsApp and Instagram—where merchants already are—and automating the painful parts of order management, Vayva offers a 10x better solution than the status quo.")
doc.add_paragraph()

add_paragraph("This marketing playbook provides the exact tactics, platforms, and strategies needed to acquire customers at low cost while building a strong brand in the African commerce ecosystem.")
doc.add_paragraph()

add_paragraph("Key Success Factors:", bold=True)
add_numbered("Consistency: Show up daily on chosen platforms")
add_numbered("Value-first: Help before selling")
add_numbered("Community: Build relationships, not just transactions")
add_numbered("Measurement: Track what matters, optimize relentlessly")
add_numbered("Iteration: Test, learn, and improve continuously")
doc.add_paragraph()

add_paragraph("The opportunity is massive. The timing is right. Execute this playbook with discipline, and Vayva will become the default commerce platform for African businesses.")
doc.add_paragraph()

# Save document
output_path = '/Users/fredrick/Documents/Vayva-Tech/vayva/Vayva_Marketing_Plan.docx'
doc.save(output_path)

print(f"Document saved to: {output_path}")
print(f"Document size: {os.path.getsize(output_path) / 1024:.1f} KB")
